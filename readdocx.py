from docx import Document
import pandas as pd

def read_docx_content_with_tables(file_path):
    doc = Document(file_path)
    # 打印doc对象
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
        # 初始化一个空字典
    print(text)
    # 获取文档中的所有表格
    tables = doc.tables
    # 遍历每个表格
    for table in tables:
        # 获取表头（第一行）
        headers = [cell.text for cell in table.rows[0].cells]
        print(headers)
        # 遍历表格的每一行（从第二行开始）
        row_data = []
        for row in table.rows[1:]:
            cell_data = []
            for i, cell in enumerate(row.cells):
                # 使用表头作为键，单元格内容作为值
                cell_data.append(cell.text)
                # row_data[headers[i]] = cell.text
            print(cell_data)
            row_data.append(cell_data)
        print(row_data)
    # print(result)
    df = pd.DataFrame(row_data,headers)
    print(df)
    # 将索引和列翻转换
    df = df.transpose()
    return df